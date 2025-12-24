# main.py
# -*- coding: utf-8 -*-
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi import Response
import traceback, logging
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
import io
import os, json, tempfile, time, re, base64
from typing import List, Dict, Any, Optional, Tuple
from urllib.parse import urlparse, urljoin
from datetime import datetime
from collections import OrderedDict
from copy import deepcopy

import httpx
import trafilatura
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("aquilas")

app = FastAPI(title="Clipping Report Builder")
STATIC_DIR = os.path.join(os.path.dirname(__file__), "static")
if os.path.isdir(STATIC_DIR):
    app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

UI_PATH = os.path.join(os.path.dirname(__file__), "ui.html")

@app.get("/", response_class=HTMLResponse)
def home():
    if os.path.exists(UI_PATH):
        return HTMLResponse(open(UI_PATH, "r", encoding="utf-8").read())
    return HTMLResponse("<h1>Clipping Report Builder</h1><p>Backend is running.</p>")

@app.head("/", include_in_schema=False)
def home_head():
    # Fast, empty response for health probes
    return Response(status_code=204)

# --- CORS ---
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# -----------------------------
# Clients (FULL NAME ONLY)
# -----------------------------
CLIENT_NAMES = [
    "ALPHA IVF GROUP BERHAD",
    "AME ELITE CONSORTIUM BERHAD",
    "AME REAL ESTATE INVESTMENT TRUST",
    "DELEUM BERHAD",
    "ECONPILE HOLDINGS BERHAD",
    "GAGASAN NADI CERGAS BERHAD",
    "GDB HOLDINGS BERHAD",
    "GUAN CHONG BERHAD",
    "KSL HOLDINGS BERHAD",
    "LAC MED BERHAD",
    "LAND & GENERAL BERHAD",
    "MALAYAN FLOUR MILLS BERHAD",
    "MALAYSIA STEEL WORKS (KL) BERHAD",
    "MATRIX CONCEPTS HOLDINGS BERHAD",
    "PECCA GROUP BERHAD",
    "SCIENTEX BERHAD",
    "SENHENG NEW RETAIL BERHAD",
    "SOUTHERN CABLE GROUP BERHAD",
    "TIONG NAM LOGISTICS HOLDINGS BERHAD",
    "TUJU SETIA BERHAD",
]

def _slug(s: str) -> str:
    out, dash = [], False
    for ch in s.lower():
        if ch.isalnum():
            out.append(ch)
            dash = False
        else:
            if not dash:
                out.append("-")
                dash = True
    return "".join(out).strip("-")

CLIENTS: List[Dict[str, str]] = [{"id": _slug(n), "name": n} for n in CLIENT_NAMES]

MEDIA_OPTIONS = [
    "Bernama", "The Star", "The Edge", "The Sun", "The Malaysian Reserve", "New Straits Times",
    "Sin Chew Daily", "China Press", "Nanyang Siang Pau", "Utusan Malaysia", "Berita Harian", "DagangNews",
]

# -----------------------------
# Models
# -----------------------------
class BuildItem(BaseModel):
    media: str
    category: str = "Online"       # "Online news" or "NewsPaper"
    section: Optional[str] = None
    language: Optional[str] = None
    date: Optional[str] = None          # ISO yyyy-mm-dd preferred
    url: Optional[str] = None           # optional for Papers
    snippet: Optional[str] = None       # optional: pre-extracted snippet
    title_override: Optional[str] = None
    image_data: Optional[str] = None    # data URL pasted from UI (e.g. "data:image/png;base64,...")

class BuildReportReq(BaseModel):
    client: str
    items: List[BuildItem]

# -----------------------------
# Helpers
# -----------------------------
def _resolve_client(client: str) -> Optional[Dict[str, str]]:
    key = (client or "").strip().lower()
    if not key:
        return None
    for c in CLIENTS:
        if key in {c["id"].lower(), c["name"].lower()}:
            return c
    return None

async def _fetch_html(url: str) -> str:
    async with httpx.AsyncClient(timeout=25, headers={"User-Agent": "Mozilla/5.0"}) as c:
        r = await c.get(url, follow_redirects=True)
        r.raise_for_status()
        return r.text

def _extract_article(url: str, html: str) -> Dict[str, Any]:
    raw = trafilatura.extract(
        html, url=url, include_images=False, include_links=False, output_format="json"
    )
    if not raw:
        return {"title": "", "text": ""}
    try:
        j = json.loads(raw)
    except Exception:
        return {"title": "", "text": ""}
    return {"title": j.get("title") or "", "text": j.get("text") or ""}

IMG_META_PATTERNS = [
    re.compile(r'<meta\s+property=["\']og:image["\']\s+content=["\']([^"\']+)["\']', re.I),
    re.compile(r'<meta\s+name=["\']twitter:image["\']\s+content=["\']([^"\']+)["\']', re.I),
    re.compile(r'<meta\s+name=["\']twitter:image:src["\']\s+content=["\']([^"\']+)["\']', re.I),
    re.compile(r'<link\s+rel=["\']image_src["\']\s+href=["\']([^"\']+)["\']', re.I),
]
IMG_FALLBACK_PAT = re.compile(r'<img\b[^>]*?\bsrc=["\']([^"\']+)["\'][^>]*>', re.I)

def _clone_para_format(src_p: Paragraph, dst_p: Paragraph):
    # style & alignment
    dst_p.style = src_p.style
    dst_p.alignment = src_p.alignment

    # paragraph-format fields
    s, d = src_p.paragraph_format, dst_p.paragraph_format
    d.left_indent = s.left_indent
    d.right_indent = s.right_indent
    d.first_line_indent = s.first_line_indent
    d.space_before = s.space_before
    d.space_after  = s.space_after
    d.line_spacing = s.line_spacing
    d.line_spacing_rule = s.line_spacing_rule
    d.keep_together = s.keep_together
    d.keep_with_next = s.keep_with_next
    d.page_break_before = s.page_break_before

async def _build_without_template(req: BuildReportReq, client_name: str) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(11)
    for idx, it in enumerate(req.items, start=1):
        _add_media_header_block(doc, client_name, it)
        art = await _fetch_article_fields(it)
        _add_article_body(doc, art["title"], art["text"], it.url, it.category or "")
        if idx < len(req.items):
            doc.add_page_break()
    return doc

# ---------- Chinese detection & font helpers ----------
CJK_RE = re.compile(r'[\u4e00-\u9fff]')
CJK_SEG_RE = re.compile(r'([\u4e00-\u9fff]+)')  # groups of CJK chars

def _has_cjk(s: Optional[str]) -> bool:
    return bool(CJK_RE.search(s or ""))

def _segment_cjk(text: str):
    """Yield (segment, is_cjk) parts preserving order."""
    parts = []
    last = 0
    text = text or ""
    for m in CJK_SEG_RE.finditer(text):
        s, e = m.span()
        if s > last:
            parts.append((text[last:s], False))
        parts.append((text[s:e], True))
        last = e
    if last < len(text):
        parts.append((text[last:], False))
    return parts

def _set_run_eastasia_font(run, ea_font: str = "SimSun", latin_font: Optional[str] = None):
    if latin_font:
        run.font.name = latin_font
    rPr = run._element.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.append(rPr)
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), ea_font)

def _force_paragraph_eastasia_font(p: Paragraph, ea_font: str = "SimSun", latin_font: Optional[str] = None):
    for r in p.runs:
        _set_run_eastasia_font(r, ea_font=ea_font, latin_font=latin_font)

def _rebuild_runs_cjk_aware(p: Paragraph, is_headline: bool, latin_font_name: str = "Trebuchet MS"):
    """Rebuild runs so CJK uses SimSun while Latin stays e.g. Trebuchet MS."""
    text_now = p.text or ""
    for r in reversed(p.runs):
        r._element.getparent().remove(r._element)
    for seg, is_cjk in _segment_cjk(text_now):
        if not seg:
            continue
        run = p.add_run(seg)
        if is_headline:
            run.bold = True
        if is_cjk:
            _set_run_eastasia_font(run, ea_font="SimSun", latin_font=None)
        else:
            run.font.name = latin_font_name  # keep Latin as template font

PARA_SPLIT_RE = re.compile(r'\n\s*\n')  # blank line or 2+ newlines
_SENT_PAT = re.compile(r'.+?(?:[.!?]["”]?|$)', re.S)  # non-greedy up to .,!,? optionally followed by ” or "

def _normalize_article_paragraphs(text: str) -> list[str]:
    if not text:
        return []

    t = text.replace("\r\n", "\n").strip()

    # 1) Preferred: split on blank lines (common for scrapers)
    if "\n\n" in t:
        blocks = [b.strip() for b in PARA_SPLIT_RE.split(t)]
        return [re.sub(r'[ \t]+', ' ', b) for b in blocks if b]

    # 2) If only single newlines exist, each line is its own paragraph
    if "\n" in t:
        lines = [ln.strip() for ln in t.split("\n")]
        paras = [ln for ln in lines if ln]
        return [re.sub(r'[ \t]+', ' ', p) for p in paras]

    # 3) Last resort: split into sentences (no lookbehind), then bundle
    sents = [m.group(0).strip() for m in _SENT_PAT.finditer(t) if m.group(0).strip()]
    if not sents:
        return [t]

    paras, buf, chars = [], [], 0
    for s in sents:
        buf.append(s)
        chars += len(s)
        # pack ~2 sentences or >400 chars per paragraph
        if len(buf) >= 2 or chars > 400:
            paras.append(" ".join(buf))
            buf, chars = [], 0
    if buf:
        paras.append(" ".join(buf))
    return paras

# ---------- Image extraction ----------
def _extract_image_url(html: str, base_url: str) -> Optional[str]:
    if not html:
        return None
    for pat in IMG_META_PATTERNS:
        m = pat.search(html)
        if m:
            return urljoin(base_url, m.group(1).strip())
    m = IMG_FALLBACK_PAT.search(html)
    if m:
        return urljoin(base_url, m.group(1).strip())
    return None

async def _download_image_to_temp(url: str) -> Optional[str]:
    if not url:
        return None
    suffix = ".jpg"
    parsed = urlparse(url)
    if "." in parsed.path.rsplit("/", 1)[-1]:
        ext = parsed.path.rsplit(".", 1)[-1].lower()
        if ext in {"jpg", "jpeg", "png", "gif", "webp"}:
            suffix = "." + ext
    try:
        async with httpx.AsyncClient(timeout=25, headers={"User-Agent": "Mozilla/5.0"}) as c:
            r = await c.get(url, follow_redirects=True)
            r.raise_for_status()
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(r.content)
                return tmp.name
    except Exception:
        return None

def _data_url_to_temp(data_url: str) -> Optional[str]:
    try:
        raw = None
        ext = ".jpg"
        if data_url.startswith("data:"):
            head, b64 = data_url.split(",", 1)
            # guess extension from mime
            if "image/png" in head:
                ext = ".png"
            elif "image/webp" in head:
                ext = ".webp"
            elif "image/gif" in head:
                ext = ".gif"
            raw = base64.b64decode(b64)
        else:
            # Assume raw base64 payload; detect type from magic bytes
            raw = base64.b64decode(data_url)
            if raw.startswith(b"\x89PNG"):
                ext = ".png"
            elif raw[:3] == b"GIF":
                ext = ".gif"
            elif raw[:4] == b"RIFF" and b"WEBP" in raw[:12]:
                ext = ".webp"
            else:
                ext = ".jpg"

        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(raw)
            return tmp.name
    except Exception:
        return None


# ---------- Fallback (no template) ----------
def _add_media_header_block(doc: Document, client_name: str, item: BuildItem):
    p = doc.add_paragraph("MEDIA CLIPPING")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(16)
    p.runs[0].bold = True

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    labels = ["Client", "Date", "Media", "Category", "Section", "Language"]
    values = [
        client_name,
        (item.date or ""),
        (item.media or ""),
        (item.category or ""),
        (item.section or ""),
        (item.language or ""),
    ]
    for i in range(6):
        row = table.rows[i]
        row.cells[0].paragraphs[0].add_run(f"{labels[i]}\t:").bold = True
        row.cells[1].paragraphs[0].add_run(values[i])
    doc.add_paragraph("")

def _add_article_body(doc: Document, title: str, text: str, url: Optional[str], category: str):
    t = doc.add_paragraph()
    run = t.add_run(title.strip() or "(Untitled)")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("")
    for para in (text or "").split("\n\n"):
        para = para.strip()
        if para:
            doc.add_paragraph(para)

    doc.add_paragraph("")
    src = doc.add_paragraph()
    if url and _is_online(category):
        src.add_run("Source from: ")
        _add_hyperlink(src, url, url)
    else:
        src.add_run("Source: Print edition").italic = True

def _rollup_summary(items: List[BuildItem], attr: str) -> str:
    vals = []
    for x in items:
        v = getattr(x, attr, None)
        if v:
            v = str(v).strip()
            if v:
                vals.append(v)
    if not vals:
        return ""
    uniq = sorted(set(vals))
    if len(uniq) == 1:
        return uniq[0]
    return f"Various ({len(uniq)})"

# ---------- Template helpers ----------
def _fmt_date(iso_str: Optional[str]) -> str:
    if not iso_str:
        return ""
    try:
        d = datetime.strptime(iso_str[:10], "%Y-%m-%d")
        return d.strftime("%d %B %Y")
    except:
        return iso_str

def _is_online(category: str) -> bool:
    cat = (category or "").lower()
    if "print" in cat or "paper" in cat:
        return False
    return True

def _replace_inline_placeholders(doc: Document, mapping: Dict[str, Any]) -> None:
    def replace_in_paragraph(p):
        original = p.text
        new_text = original
        for k, v in mapping.items():
            token = f"{{{{{k}}}}}"
            if token in new_text:
                new_text = new_text.replace(token, str(v))
        if new_text != original:
            for r in reversed(p.runs):
                r._element.getparent().remove(r._element)
            p.add_run(new_text)

    for p in doc.paragraphs:
        replace_in_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

    for section in doc.sections:
        for hf in (section.header, section.footer):
            for p in hf.paragraphs:
                replace_in_paragraph(p)
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_in_paragraph(p)

def _find_paragraph_with_text(doc: Document, needle: str):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    return None

def _find_last_paragraph_with_text(doc: Document, needle: str):
    for p in reversed(doc.paragraphs):
        if needle in p.text:
            return p
    return None

def _iter_block_items(parent_doc: Document):
    body = parent_doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent_doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent_doc)

# ---------- Hyperlink ----------
def _add_hyperlink(paragraph, url: str, text: Optional[str] = None):
    if not url:
        return paragraph.add_run("")
    if text is None:
        text = url
    try:
        part = paragraph.part
        r_id = part.relate_to(url, reltype=RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rPr.append(color)
        new_run.append(rPr)
        new_text = OxmlElement('w:t')
        new_text.text = text
        new_run.append(new_text)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return paragraph
    except Exception:
        run = paragraph.add_run(text)
        run.underline = True
        return paragraph

# ---------- Article fetch (title, text, image-url + pasted image) ----------
async def _fetch_article_fields(item: BuildItem) -> Dict[str, Any]:
    title = item.title_override or ""
    text = item.snippet or ""
    image_url = None
    pasted_image_path = None

    if item.image_data:
        pasted_image_path = _data_url_to_temp(item.image_data)
        if pasted_image_path:
            logger.info(f"[newspaper image] got pasted image ({len(item.image_data)} chars) -> {pasted_image_path}")
        else:
            logger.warning("[newspaper image] image_data present but could not decode (check data URL/base64)")


    is_newspaper = not _is_online(item.category or "")

    if is_newspaper:
        # Newspaper: DO NOT fetch headline/content/URL or image from the web.
        # We rely solely on user-pasted image (image_data).
        pass
    else:
        # Online: may fetch headline/content/image
        if item.url:
            try:
                html = await _fetch_html(item.url)
                if not title or not text:
                    meta = _extract_article(item.url, html)
                    title = title or (meta.get("title") or "")
                    text = text or (meta.get("text") or "")
                if not pasted_image_path:
                    image_url = _extract_image_url(html, item.url)
            except Exception:
                pass

    # Placeholders only for ONLINE path (newspaper mapping blanks these anyway)
    if not is_newspaper:
        if not text or not text.strip():
            text = "(Content not extracted. See source note below.)"
            if not title or title.strip() == "":
                title = "[Headline]"
        else:
            if not title or title.strip() == "":
                if item.url:
                    title = urlparse(item.url).path.rsplit("/", 1)[-1].replace("-", " ").title() or "[Headline]"
                else:
                    title = "[Headline]"

    url_label = ""
    if item.media and item.media.strip():
        url_label = item.media.strip()
    elif item.url and not is_newspaper:
        host = urlparse(item.url).hostname or ""
        url_label = host.replace("www.", "") if host else "link"

    return {
        "title": title,
        "text": text,
        "image_url": image_url,
        "pasted_image_path": pasted_image_path,
        "url_label": url_label,
    }




# ---------- Card prototype detection & cloning ----------
ITEM_TOKENS = {
    "{{ITEM_HEADLINE}}",
    "{{ITEM_IMAGE}}",
    "{{ITEM_CONTENT}}",
    "{{ITEM_URL}}",
    "{{ITEM_MEDIA}}",
    "{{ITEM_SECTION}}",
    "{{ITEM_LANGUAGE}}",
    "{{ITEM_DATE}}",
    "{{ITEM_URL_LABEL}}",
}

def _block_has_item_tokens(block: Any) -> bool:
    if isinstance(block, Paragraph):
        t = block.text or ""
        return any(tok in t for tok in ITEM_TOKENS)
    if isinstance(block, Table):
        for row in block.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text or ""
                    if any(tok in t for tok in ITEM_TOKENS):
                        return True
    return False

def _find_extracts_marker_and_card(doc: Document) -> Tuple[Optional[Paragraph], List[Any]]:
    marker = None
    blocks = list(_iter_block_items(doc))
    for i, b in enumerate(blocks):
        if isinstance(b, Paragraph) and "{{EXTRACTS_START}}" in (b.text or ""):
            marker = b
            card_blocks = []
            j = i + 1
            while j < len(blocks) and _block_has_item_tokens(blocks[j]):
                card_blocks.append(blocks[j])
                j += 1
            return marker, card_blocks
    return None, []

def _remove_blocks(blocks: List[Any]) -> None:
    for b in blocks:
        el = b._element
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

def _insert_blocks_after(ref_para_or_tbl: Any, blocks_xml: List[Any]) -> List[Any]:
    inserted = []
    ref_el = ref_para_or_tbl._element
    for el in blocks_xml:
        new_el = deepcopy(el)
        ref_el.addnext(new_el)
        ref_el = new_el
        inserted.append(new_el)
    return inserted

def _replace_in_paragraph_single(p: Paragraph, mapping: Dict[str, Any],
                                 url_for_link: Optional[str], url_label: str) -> None:
    # Snapshot the original text/tokens before any mutation
    original_text   = p.text or ""
    had_headline    = "{{ITEM_HEADLINE}}" in original_text
    had_url_token   = "{{ITEM_URL}}" in original_text
    has_content_tok = "{{ITEM_CONTENT}}" in original_text

    # ---------- 1) Expand {{ITEM_CONTENT}} into real paragraphs ----------
    if has_content_tok:
        content_text = str(mapping.get("ITEM_CONTENT", "") or "")
        content_paras = _normalize_article_paragraphs(content_text)  # returns list[str]

        # Always have at least one paragraph so .replace() works cleanly
        if not content_paras:
            content_paras = [""]

        # Replace the token in the current paragraph with the FIRST content paragraph
        first_text = content_paras[0]
        base_text = original_text.replace("{{ITEM_CONTENT}}", first_text)

        # Clear runs and write the updated base text back
        for r in list(p.runs)[::-1]:
            r._element.getparent().remove(r._element)
        p.add_run(base_text if base_text else "")

        # Insert remaining content paragraphs AFTER this paragraph
        # Use the official underlying element handle (_p) when adding siblings
        after_el = p._p
        parent_container = p._parent  # works for body, header/footer, table cell

        for extra in content_paras[1:]:
            new_el = OxmlElement('w:p')
            after_el.addnext(new_el)
            after_el = new_el
            new_p = Paragraph(new_el, parent_container)
            _clone_para_format(p, new_p)   # inherit spacing/alignment/style
            if extra:
                new_p.add_run(extra)

        # Refresh our working text for subsequent simple replacements
        current_text = p.text or ""
    else:
        current_text = original_text

    # ---------- 2) Replace simple tokens (excluding IMAGE/URL/CONTENT) ----------
    for k, v in mapping.items():
        if k in ("ITEM_IMAGE", "ITEM_URL", "ITEM_CONTENT"):
            continue
        token = f"{{{{{k}}}}}"
        if token in current_text:
            current_text = current_text.replace(token, str(v))

    if current_text != (p.text or ""):
        for r in list(p.runs)[::-1]:
            r._element.getparent().remove(r._element)
        p.add_run(current_text if current_text else "")

    # ---------- 3) URL token -> hyperlink line ----------
    if had_url_token:
        # Clear all current runs first
        for r in list(p.runs)[::-1]:
            r._element.getparent().remove(r._element)

        if url_for_link:
            # REQUIRED: no media label — just "Source from: {url}"
            p.add_run("Source from: ")
            _add_hyperlink(p, url_for_link, url_for_link)
        else:
            # No URL provided: remove the token entirely
            # (prevents stray "{{ITEM_URL}}" appearing in output)
            pass

    # ---------- 4) Headline styling ----------
    if had_headline:
        for r in p.runs:
            r.bold = True


def _insert_image_into_paragraph(p: Paragraph, image_path: Optional[str], keep_original: bool = False):
    for r in reversed(p.runs):
        r._element.getparent().remove(r._element)
    if image_path:
        try:
            run = p.add_run()
            if keep_original:
                # original native size
                run.add_picture(image_path)
            else:
                # fixed width for online items
                run.add_picture(image_path, width=Cm(10))
            return
        except Exception:
            pass
    p.add_run("")


def _replace_placeholders_in_inserted_elements(
    doc: Document,
    inserted_elements: List[Any],
    mapping: Dict[str, Any],
    image_path: Optional[str],
    url_for_link: Optional[str],
    url_label: str,
    use_chinese_font: bool = False,
    keep_original_image: bool = False,
) -> None:
    """
    IMPORTANT: Detect {{ITEM_IMAGE}} BEFORE replacing text, then insert image afterwards.
    If use_chinese_font=True, rebuild runs for headline/content so Chinese uses SimSun
    and Latin stays Trebuchet MS (URL lines are left as template-styled).
    """

    # NEW: track whether any image token existed in this card
    found_image_token = False

    for el in inserted_elements:
        if isinstance(el, CT_P):
            p = Paragraph(el, doc)
            original_text = p.text or ""
            had_img = "{{ITEM_IMAGE}}" in original_text
            had_headline = "{{ITEM_HEADLINE}}" in original_text
            had_url_token = "{{ITEM_URL}}" in original_text

            _replace_in_paragraph_single(p, mapping, url_for_link, url_label)

            if had_img:
                found_image_token = True
                _insert_image_into_paragraph(p, image_path, keep_original=keep_original_image)
            else:
                if use_chinese_font and not had_url_token:
                    _rebuild_runs_cjk_aware(p, is_headline=had_headline)

        elif isinstance(el, CT_Tbl):
            tbl = Table(el, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        original_text = p.text or ""
                        had_img = "{{ITEM_IMAGE}}" in original_text
                        had_headline = "{{ITEM_HEADLINE}}" in original_text
                        had_url_token = "{{ITEM_URL}}" in original_text

                        _replace_in_paragraph_single(p, mapping, url_for_link, url_label)

                        if not had_img:  # We don't bold the image itself
                            for run in p.runs:
                                run.bold = True
                                
                        if had_img:
                            found_image_token = True
                            _insert_image_into_paragraph(p, image_path, keep_original=keep_original_image)
                        else:
                            if use_chinese_font and not had_url_token:
                                _rebuild_runs_cjk_aware(p, is_headline=had_headline)

    # Fallback: if there was a pasted image but NO {{ITEM_IMAGE}} anywhere,
    # insert it into the first block of the card. Prefer the first table cell, else a paragraph before the card.
    if image_path and not found_image_token and inserted_elements:
        first_el = inserted_elements[0]

        # Case A: the card starts with a table -> put image into the top-left cell
        if isinstance(first_el, CT_Tbl):
            tbl = Table(first_el, doc)
            # top-left cell
            cell = tbl.rows[0].cells[0]
            # use the first paragraph in that cell (or create one)
            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            # clear current runs and add the picture
            for r in list(p.runs)[::-1]:
                r._element.getparent().remove(r._element)
            run = p.add_run()
            if keep_original_image:
                run.add_picture(image_path)
            else:
                run.add_picture(image_path, width=Cm(10))

        # Case B: the card starts with a paragraph -> insert a paragraph *before* the card and add the image
        else:
            new_el = OxmlElement('w:p')
            first_el.addprevious(new_el)
            new_p = Paragraph(new_el, doc)
            run = new_p.add_run()
            if keep_original_image:
                run.add_picture(image_path)
            else:
                run.add_picture(image_path, width=Cm(10))



# ---------- Build using template (per-media pages, no clippings table) ----------
async def _build_with_template(template_path: str, payload: BuildReportReq, client_name: str) -> Document:
    items: List[BuildItem] = payload.items or []

    # Group by media + category(normalized to online/newspaper) + language, preserving order
    def _cat_key(cat: str) -> str:
        return "online" if _is_online(cat) else "print"

    media_groups: "OrderedDict[tuple[str, str, str], List[BuildItem]]" = OrderedDict()
    for it in items:
        key = (
            (it.media or "").strip(),
            _cat_key(it.category or ""),
            (it.language or "").strip().lower(),
        )
        media_groups.setdefault(key, []).append(it)

    total = len(items)
    online = sum(1 for x in items if _is_online(x.category))
    paper = total - online

    doc = Document(template_path)

    # Find header & card prototype on first page
    marker_para, card_blocks_in_doc = _find_extracts_marker_and_card(doc)
    if not marker_para:
        marker_para = doc.add_paragraph("")
    card_proto_xml = [deepcopy(b._element) for b in card_blocks_in_doc] if card_blocks_in_doc else []
    _remove_blocks(card_blocks_in_doc)
    marker_para.text = ""

    # FIRST MEDIA
    (first_media, first_cat_key, first_lang), first_items = next(iter(media_groups.items())) if media_groups else (("", "", ""), [])
    dates_first = sorted([x.date for x in first_items if x.date])
    date_single_first = _fmt_date(dates_first[0]) if dates_first else ""

    mapping_first = {
        "REPORT_TITLE": "Media Coverage Report",
        "CLIENT_NAME": client_name,
        "DATE_RANGE": date_single_first,  # single date
        "TOTAL_CLIPPINGS": total,
        "ONLINE_COUNT": online,
        "NEWSPAPER_COUNT": paper,
        "SUMMARY_TEXT": f"Total clippings: {total}. Online: {online}. Newspaper: {paper}.",
        "LANG_SUMMARY": _rollup_summary(first_items, "language"),  # stays single now
        "MEDIA_SUMMARY": first_media or "",
        "CATEGORY_SUMMARY": "Online" if first_cat_key == "online" else "Print",
        "SECTION_SUMMARY": _rollup_summary(first_items, "section"),
    }
    _replace_inline_placeholders(doc, mapping_first)

    # Clear any clippings marker on first page
    anchor = _find_paragraph_with_text(doc, "{{CLIPPINGS_TABLE}}")
    if anchor is not None:
        anchor.text = ""

    # Render items for first media
    insertion_ref = marker_para
    if card_proto_xml:
        for it in first_items:
            art = await _fetch_article_fields(it)
            is_newspaper = not _is_online(it.category or "")
            mapping_item = {
                "ITEM_HEADLINE": art["title"] if not is_newspaper else "",
                "ITEM_CONTENT": art["text"]   if not is_newspaper else "",
                "ITEM_URL":     (it.url or "") if (it.url and not is_newspaper) else "",
                "ITEM_URL_LABEL": art["url_label"],
                "ITEM_MEDIA": it.media or "",
                "ITEM_SECTION": it.section or "",
                "ITEM_LANGUAGE": it.language or "",
                "ITEM_DATE": _fmt_date(it.date),
                "ITEM_IMAGE": "",  # handled separately
            }

            image_path = None
            if art["pasted_image_path"]:
                image_path = art["pasted_image_path"]
            elif art.get("image_url"):
                image_path = await _download_image_to_temp(art["image_url"])

            use_chinese_font = (it.language or "").strip().lower() == "chinese"

            inserted = _insert_blocks_after(insertion_ref, card_proto_xml)
            _replace_placeholders_in_inserted_elements(
                doc,
                inserted,
                mapping_item,
                image_path,
                (it.url if (it.url and not is_newspaper) else None),
                art["url_label"],
                use_chinese_font=use_chinese_font,
                keep_original_image=is_newspaper,
            )
            insertion_ref = Paragraph(inserted[-1], doc) if inserted else insertion_ref
    else:
        for i, it in enumerate(first_items, start=1):
            art = await _fetch_article_fields(it)
            tpara = doc.add_paragraph()
            run = tpara.add_run(f"{i}. {art['title']}")
            run.bold = True
            run.font.size = Pt(12)
            meta_line = f"{it.media or ''} • {it.section or ''} • {it.language or ''} • {_fmt_date(it.date)}"
            doc.add_paragraph(meta_line).italic = True
            if art['text']:
                doc.add_paragraph(art['text'])
            if it.url and _is_online(it.category or ""):
                p = doc.add_paragraph()
                p.add_run("Source from: ")
                _add_hyperlink(p, it.url, it.url)
            doc.add_paragraph().add_run("—" * 40)

    # Remaining media pages
    def _capture_header_blocks(path: str) -> List[Any]:
        proto = Document(path)
        blocks_xml = []
        for block in _iter_block_items(proto):
            blocks_xml.append(block._element)
            if isinstance(block, Paragraph) and "{{EXTRACTS_START}}" in block.text:
                break
        return [deepcopy(el) for el in blocks_xml]

    header_proto_xml = _capture_header_blocks(template_path)

    remaining = list(media_groups.items())[1:]
    for (media_name, cat_key, lang_key), group_items in remaining:
        doc.add_page_break()
        end_anchor = doc.add_paragraph("")
        inserted_hdr = _insert_blocks_after(end_anchor, header_proto_xml)
        if inserted_hdr:
            parent = end_anchor._element.getparent()
            if parent is not None:
                parent.remove(end_anchor._element)

        dates_group = sorted([x.date for x in group_items if x.date])
        date_single = _fmt_date(dates_group[0]) if dates_group else ""
        mapping_group = {
            "REPORT_TITLE": "Media Coverage Report",
            "CLIENT_NAME": client_name,
            "DATE_RANGE": date_single,
            "TOTAL_CLIPPINGS": total,
            "ONLINE_COUNT": online,
            "NEWSPAPER_COUNT": paper,
            "SUMMARY_TEXT": f"Total clippings: {total}. Online: {online}. Newspaper: {paper}.",
            "LANG_SUMMARY": _rollup_summary(group_items, "language"),
            "MEDIA_SUMMARY": media_name or "",
            "CATEGORY_SUMMARY": "Online news" if cat_key == "online" else "Newspaper",
            "SECTION_SUMMARY": _rollup_summary(group_items, "section"),
        }
        _replace_inline_placeholders(doc, mapping_group)

        last_table_marker = _find_last_paragraph_with_text(doc, "{{CLIPPINGS_TABLE}}")
        if last_table_marker is not None:
            last_table_marker.text = ""

        last_extracts_marker = _find_last_paragraph_with_text(doc, "{{EXTRACTS_START}}")
        if last_extracts_marker:
            last_extracts_marker.text = ""
            insertion_ref = last_extracts_marker
        else:
            insertion_ref = doc.add_paragraph("")

        if card_proto_xml:
            for it in group_items:
                art = await _fetch_article_fields(it)
                is_newspaper = not _is_online(it.category or "")

                mapping_item = {
                    "ITEM_HEADLINE": art["title"] if not is_newspaper else "",
                    "ITEM_CONTENT": art["text"]   if not is_newspaper else "",
                    "ITEM_URL":     (it.url or "") if (it.url and not is_newspaper) else "",
                    "ITEM_URL_LABEL": art["url_label"],
                    "ITEM_MEDIA": it.media or "",
                    "ITEM_SECTION": it.section or "",
                    "ITEM_LANGUAGE": it.language or "",
                    "ITEM_DATE": _fmt_date(it.date),
                    "ITEM_IMAGE": "",
                }

                image_path = None
                if art["pasted_image_path"]:
                    image_path = art["pasted_image_path"]
                elif art.get("image_url"):
                    image_path = await _download_image_to_temp(art["image_url"])

                use_chinese_font = (it.language or "").strip().lower() == "chinese"

                inserted = _insert_blocks_after(insertion_ref, card_proto_xml)
                _replace_placeholders_in_inserted_elements(
                    doc,
                    inserted,
                    mapping_item,
                    image_path,
                    (it.url if (it.url and not is_newspaper) else None),
                    art["url_label"],
                    use_chinese_font=use_chinese_font,
                    keep_original_image=is_newspaper,
                )
                insertion_ref = Paragraph(inserted[-1], doc) if inserted else insertion_ref
        else:
            for i, it in enumerate(group_items, start=1):
                art = await _fetch_article_fields(it)
                tpara = doc.add_paragraph()
                run = tpara.add_run(f"{i}. {art['title']}")
                run.bold = True
                run.font.size = Pt(12)
                meta_line = f"{it.media or ''} • {it.section or ''} • {it.language or ''} • {_fmt_date(it.date)}"
                doc.add_paragraph(meta_line).italic = True
                if art['text']:
                    doc.add_paragraph(art['text'])
                if it.url and _is_online(it.category or ""):
                    p = doc.add_paragraph()
                    p.add_run("Source from: ")
                    _add_hyperlink(p, it.url, it.url)
                doc.add_paragraph().add_run("—" * 40)

    return doc


# -----------------------------
# API
# -----------------------------
@app.get("/healthz")
def healthz():
    return {"status": "ok"}

@app.get("/clients")
def clients():
    return CLIENTS

@app.get("/media")
def media():
    return MEDIA_OPTIONS

@app.get("/debug/template")
def debug_template():
    path = os.path.join(os.path.dirname(__file__), "template.docx")
    exists = os.path.exists(path)
    info = {
        "template_expected_path": path,
        "exists": exists,
    }
    if exists:
        stat = os.stat(path)
        info.update({
            "size_bytes": stat.st_size,
            "last_modified_iso": time.strftime("%Y-%m-%dT%H:%M:%S", time.localtime(stat.st_mtime)),
        })
    return info

@app.post("/build_report")
async def build_report(req: BuildReportReq):
    try:
        if not req.items:
            return JSONResponse(status_code=400, content={"error": "No items in request."})

        c = _resolve_client(req.client) or {"name": req.client}
        client_name = c["name"]

        template_path = os.path.join(os.path.dirname(__file__), "template.docx")
        use_template = os.path.exists(template_path)

        if use_template:
            try:
                doc = await _build_with_template(template_path, req, client_name)
            except Exception:
                logger.exception("Template build failed; falling back to no-template layout.")
                doc = await _build_without_template(req, client_name)
                use_template = False
        else:
            doc = await _build_without_template(req, client_name)

        # --- Save to temp, then read bytes and return (avoid streaming issues) ---
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_client = (client_name or "client").replace(" ", "_")
        filename = f"media_clipping_{safe_client}_{stamp}.docx"

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp_path = tmp.name

        doc.save(tmp_path)

        with open(tmp_path, "rb") as f:
            data = f.read()

        # Optionally remove the temp file (not required on Render’s ephemeral FS)
        try:
            os.remove(tmp_path)
        except Exception:
            pass

        extra_headers = {
            "X-Template-Used": "1" if use_template else "0",
            "Content-Disposition": f'attachment; filename="{filename}"',
        }
        if use_template:
            extra_headers["X-Template-Path"] = template_path

        return Response(
            content=data,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=extra_headers,
        )

    except Exception as e:
        logger.exception("build_report crashed")
        return JSONResponse(
            status_code=500,
            content={"error": str(e), "trace": traceback.format_exc()},
        )











